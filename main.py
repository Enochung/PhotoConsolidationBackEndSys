# 匯入 Flask 框架的核心模組
# Flask: 用於創建 Web 應用的核心類
# request: 用於處理 HTTP 請求數據
# jsonify: 將 Python 對象轉換為 JSON 格式響應
from flask import Flask, request, jsonify
# CORS（跨域資源共享）允許前端和後端在不同域名下進行通信
from flask_cors import CORS
# python-docx 庫，用於生成和操作 Word 文件（.docx 格式）
from docx import Document
# os 模組，用於操作文件和目錄，例如新增、刪除文件夾
import os
# time 模組，用於生成當前時間戳，通常用於創建唯一文件名
import time
# send_file 用於將文件作為附件發送給客戶端
from flask import send_file

# 初始化 Flask 應用程式
# __name__ 是 Python 預定義變量，用於告知 Flask 當前模組名稱
app = Flask(__name__)
# 啟用 CORS，允許來自 http://localhost:4200 的跨域請求
CORS(app, resources={r"/*": {"origins": "http://localhost:4200"}})
# 定義上傳目錄，所有文件將儲存在 'uploads' 文件夾中
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
# 如果目錄不存在，則新增該目錄
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 定義處理圖片和說明上傳的 API 路由
@app.route('/flaskapi/api/upload', methods=['POST'])
def upload_file():
    try:
        # 確認請求中是否包含 'images' 文件
        if 'images' not in request.files:
            return jsonify({"error": "未上傳圖片"}), 400  # 如果沒有上傳圖片，回應錯誤

        # 獲取圖片列表和說明（如果未提供值，設置為預設值）
        image_files = request.files.getlist('images')  # 獲取上傳的圖片列表
        title = request.form.get('title', '').strip() or '預設標題'  # 標題
        description = request.form.get('description', '').strip() or '無提供說明'  # 描述
        shooting_time = request.form.get('shooting_time', '').strip() or f'{time.strftime("%Y%m%d")}'  # 拍攝時間
        shooting_location = request.form.get('shooting_location', '').strip() or '台灣'  # 拍攝地點
        photographer = request.form.get('photographer', '').strip() or 'None'  # 攝影人

        # 新增 Word 文件並插入標題和說明
        document = Document()  # 新增一個新的 Word 檔案的物件
        document.add_heading('圖片報告', 0)  # 添加文件主標題
        document.add_paragraph(f'標題: {title}')  # 插入描述段落標題

        # forloop上傳的圖片文件，將其插入 Word 文件表格
        for index, image_file in enumerate(image_files):
            # 構建圖片的完整文件路徑並保存
            image_path = os.path.join(UPLOAD_FOLDER, image_file.filename)
            image_file.save(image_path)  # 將文件保存到指定目錄

            # 新增表格，插入圖片和描述信息
            table = document.add_table(rows=3, cols=6)  #  3 row 和 6 columns 的表格
            table.style = 'Table Grid'  # 設置表格樣式為網格

            # 第一行，合併所有列來插入圖片
            row1 = table.rows[0]  # 獲取表格的第一行（準備在該行插入圖片）
            cell1 = row1.cells[0].merge(row1.cells[5])  # 合併第一行的所有單元格（從第1列到第6列），用於插入圖片
            paragraph = cell1.paragraphs[0]  # 獲取合併後單元格中的段落（默認會有一個空段落）
            run = paragraph.add_run()  # 為段落添加新的運行（Run），用於插入圖片
            run.add_picture(image_path, width=5000000)  # 在段落中插入圖片，指定圖片的寬度（5000000 EMUs，約為5英寸）

            # 第二行，填寫表格資訊
            table.cell(1, 0).text = '照片編號'  # 在第二行第一列的單元格中填寫標題文字 "照片編號"
            table.cell(1, 1).text = f'{index + 1:02d}'  # 在第二行第二列的單元格中填寫照片編號，編號格式為兩位數（如 01, 02）
            table.cell(1, 2).text = '說明'  # 在第二行第三列的單元格中填寫標題文字 "說明"
            row2 = table.rows[1]  # 獲取表格的第二行（索引為 1）
            cell2 = row2.cells[3].merge(row2.cells[5])  # 合併第二行中第4列至第6列的單元格，形成一個大單元格
            cell2.text = f'{description}'  # 將提供的描述文字插入到合併後的單元格中

            # 第三行，攝影時間、地點和人員資訊
            table.cell(2, 0).text = '攝影時間'
            table.cell(2, 1).text = f'{shooting_time}'  # 攝影時間
            table.cell(2, 2).text = '攝影地點'
            table.cell(2, 3).text = f'{shooting_location}'  # 攝影地點
            table.cell(2, 4).text = '攝影人'
            table.cell(2, 5).text = f'{photographer}'  # 攝影人

            # 每插入兩個表格後，添加一個段落作為標題
            if (index + 1) % 2 == 0 and (index + 1) != len(image_files):
                document.add_page_break()
                document.add_paragraph(f'標題: {title}')
        # 保存 Word 文件
        word_file_path = os.path.join(UPLOAD_FOLDER, f'{title}_{int(time.time())}.docx')
        # 將 Word 文件保存到伺服器的目錄中
        document.save(word_file_path)
        # 刪除已保存的圖片文件
        delete_file_in_folder(UPLOAD_FOLDER)
        # 返回response，包括生成的 Word 文件路徑
        return jsonify({"message": "文件已成功處理", "word_file": word_file_path}), 200
    except Exception as e:
        # 如果發生任何異常，返回錯誤信息
        return jsonify({"error": str(e)}), 500

# 刪除文件的function
def delete_file_in_folder(folder_path, specific_file=None):
    try:
        # 如果提供了特定文件名稱
        if specific_file:
            # 拼接完整檔案路徑
            file_path = os.path.join(folder_path, specific_file)
            # 確保檔案存在
            if os.path.exists(file_path):
                # 刪除檔案
                os.remove(file_path)
                # 返回成功訊息
                return {"message": f"成功刪除指定檔案: {specific_file}"}
            else:
                # 返回檔案不存在的錯誤消息
                return {"error": f"檔案 {specific_file} 不存在"}
        # 未指定文件名稱，刪除所有圖片檔案
        else:
            # 用於記錄已刪除的文件
            deleted_files = []
            # forloop目錄中的所有文件
            for file_name in os.listdir(folder_path):
                # 拼接完整文件路徑
                file_path = os.path.join(folder_path, file_name)
                # 判斷是否為圖片文件
                if file_name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                    # 刪除圖片文件
                    os.remove(file_path)
                    # 記錄刪除的文件名稱
                    deleted_files.append(file_name)
            # 返回刪除結果
            if deleted_files:
                return {"message": "成功刪除所有圖片檔案", "deleted_files": deleted_files}
            else:
                # 返回沒有檔案需要刪除的錯誤訊息
                return {"message": "沒有圖片檔案需要刪除"}
    except Exception as e:
        # 如果發生任何異常，返回錯誤信息
        return {"error": f"刪除圖片文件時出現錯誤: {str(e)}"}

# 定義處理文件瀏覽的 API 路由
@app.route('/flaskapi/api/files', methods=['GET'])
def list_docx_files():
    try:
        # 獲取目錄中所有以 .docx 結尾的檔案
        docx_files = [f for f in os.listdir(UPLOAD_FOLDER) if f.endswith('.docx')]
        # 返回文件列表
        return jsonify({"docx_files": docx_files}), 200
    except FileNotFoundError:
        # 如果目錄不存在，返回 404 錯誤
        return jsonify({"error": "文件未找到"}), 404
    except Exception as e:
        # 捕獲異常，返回 500 錯誤和錯誤訊息
        return jsonify({"error": str(e)}), 500

# 定義處理文件下載的 API 路由
@app.route('/flaskapi/api/download/<filename>', methods=['POST'])
def download_file(filename):
    try:
        # 拼接文件的完整路徑
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        # 判斷檔案是否存在
        if not os.path.exists(file_path):
            # 如果檔案不存在，返回 404 錯誤
            return jsonify({"error": "檔案不存在"}), 404
        # 發送文件，作為附件下載
        return send_file(file_path, as_attachment=True)
    except FileNotFoundError:
        # 如果檔案未找到，返回 404 錯誤
        return jsonify({"error": "檔案未找到"}), 404
    except Exception as e:
        # 捕獲異常，返回 500 錯誤和錯誤訊息
        return jsonify({"error": str(e)}), 500

# 定義處理文件刪除的 API 路由
@app.route('/flaskapi/api/delete/<filename>', methods=['POST'])
def delete_file(filename):
    try:
        # 呼叫 delete_file_in_folder function 執行刪除操作
        result = delete_file_in_folder(UPLOAD_FOLDER, filename)
        # 判斷是否有錯誤
        if "error" in result:
            if "刪除圖片文件時出現錯誤" in result["error"]:
                # 返回 500 錯誤，表示伺服器內部錯誤
                return jsonify(result), 500
            else:
                # 返回 404 錯誤，表示檔案未找到
                return jsonify(result), 404
        # 如果刪除成功，返回成功訊息
        return jsonify(result), 200
    except FileNotFoundError:
        # 如果檔案未找到，返回 404 錯誤
        return jsonify({"error": "檔案未找到"}), 404
    except Exception as e:
        # 捕獲異常，返回 500 錯誤和錯誤訊息
        return jsonify({"error": str(e)}), 500

# 啟動 Flask 應用
if __name__ == '__main__':
    # 啟用debug模式
    app.run(debug=True)
